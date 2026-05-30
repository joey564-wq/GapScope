"""Export a structured Resume to a Word (.docx) document, in memory.

Returns bytes (a .docx is a zip archive) so the API layer can stream it as a
download without ever writing resume content to disk — consistent with the
no-storage privacy posture. The full resume, including name and contact details,
is rendered here locally; this data never went to the LLM.
"""

from __future__ import annotations

import io

from docx import Document

from .models import Resume


def resume_to_docx_bytes(resume: Resume) -> bytes:
    """Render a Resume to .docx and return the file as bytes."""
    doc = Document()

    doc.add_heading(resume.name, level=0)

    contact_bits = [
        resume.contact.location,
        resume.contact.phone,
        resume.contact.email,
        *resume.contact.links,
    ]
    contact_line = " | ".join(b for b in contact_bits if b)
    if contact_line:
        doc.add_paragraph(contact_line)

    if resume.summary:
        doc.add_heading("Summary", level=1)
        doc.add_paragraph(resume.summary)

    for entry in resume.entries:
        header = entry.title
        if entry.organization:
            header += f" — {entry.organization}"
        doc.add_heading(header, level=1)

        meta_bits = [entry.location, entry.date_range]
        meta_line = " | ".join(b for b in meta_bits if b)
        if meta_line:
            doc.add_paragraph(meta_line, style="Intense Quote")

        for bullet in entry.bullets:
            doc.add_paragraph(bullet.text, style="List Bullet")

    if resume.skill_groups:
        doc.add_heading("Skills", level=1)
        for group in resume.skill_groups:
            doc.add_paragraph(f"{group.category}: {', '.join(group.skills)}")

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
