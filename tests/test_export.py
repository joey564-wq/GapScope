"""Tests for docx export."""

from __future__ import annotations

import io

from docx import Document

from resume_tailor import Resume
from resume_tailor.export import resume_to_docx_bytes


def test_export_returns_a_docx_zip(resume: Resume) -> None:
    data = resume_to_docx_bytes(resume)
    assert isinstance(data, bytes)
    # A .docx is a zip archive; zip files start with the bytes "PK".
    assert data[:2] == b"PK"


def test_export_contains_name_and_bullets(resume: Resume) -> None:
    data = resume_to_docx_bytes(resume)
    doc = Document(io.BytesIO(data))
    text = "\n".join(p.text for p in doc.paragraphs)
    headings = "\n".join(p.text for p in doc.paragraphs)
    assert resume.name in text or resume.name in headings
    # A known bullet from the fixture should be present.
    assert any("service-health monitor" in p.text for p in doc.paragraphs)
